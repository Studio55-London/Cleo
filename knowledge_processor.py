"""
Knowledge Base Processing Module
Handles document extraction, chunking, embedding, and entity extraction for GraphRAG
"""
import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional

# Document processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Token counting
try:
    import tiktoken
except ImportError:
    tiktoken = None


class KnowledgeProcessor:
    """Process documents for knowledge base"""

    def __init__(self, chunk_size=512, chunk_overlap=50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = self._init_tokenizer()

    def _init_tokenizer(self):
        """Initialize tiktoken tokenizer"""
        if tiktoken:
            try:
                return tiktoken.get_encoding("cl100k_base")
            except:
                pass
        return None

    # ===================================
    # Document Extraction
    # ===================================

    def extract_text(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from document based on file type
        Returns: (text_content, metadata)
        """
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()

        if file_ext == '.pdf':
            return self.extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return self.extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    def extract_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from PDF"""
        if not PyPDF2:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

        text = []
        metadata = {'pages': 0, 'extracted_method': 'PyPDF2'}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)

            return '\n\n'.join(text), metadata
        except Exception as e:
            raise Exception(f"Failed to extract PDF: {str(e)}")

    def extract_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from DOCX"""
        if not DocxDocument:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        try:
            doc = DocxDocument(file_path)
            text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            metadata = {
                'paragraphs': len(doc.paragraphs),
                'extracted_method': 'python-docx'
            }

            return '\n\n'.join(text), metadata
        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")

    def extract_txt(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from TXT/MD files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

            metadata = {
                'lines': len(text.split('\n')),
                'extracted_method': 'plain_text'
            }

            return text, metadata
        except Exception as e:
            raise Exception(f"Failed to extract text file: {str(e)}")

    # ===================================
    # Text Chunking
    # ===================================

    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """
        Split text into chunks with overlap
        Returns list of chunk dictionaries with metadata
        """
        if not text:
            return []

        # Clean text
        text = self._clean_text(text)

        # Use token-based chunking if tokenizer available
        if self.tokenizer:
            return self._chunk_by_tokens(text, metadata)
        else:
            return self._chunk_by_chars(text, metadata)

    def _chunk_by_tokens(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Chunk text by token count"""
        tokens = self.tokenizer.encode(text)
        chunks = []

        start = 0
        chunk_index = 0

        while start < len(tokens):
            end = start + self.chunk_size
            chunk_tokens = tokens[start:end]
            chunk_text = self.tokenizer.decode(chunk_tokens)

            chunks.append({
                'chunk_index': chunk_index,
                'content': chunk_text.strip(),
                'token_count': len(chunk_tokens),
                'metadata': metadata or {}
            })

            start = end - self.chunk_overlap
            chunk_index += 1

        return chunks

    def _chunk_by_chars(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Chunk text by character count (fallback)"""
        # Approximate tokens (rough estimate: 1 token ≈ 4 chars)
        char_chunk_size = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + char_chunk_size
            chunk_text = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk_text.rfind('. ')
                if last_period > char_chunk_size * 0.7:  # At least 70% of chunk size
                    chunk_text = chunk_text[:last_period + 1]
                    end = start + last_period + 1

            chunks.append({
                'chunk_index': chunk_index,
                'content': chunk_text.strip(),
                'token_count': self._estimate_tokens(chunk_text),
                'metadata': metadata or {}
            })

            start = end - char_overlap
            chunk_index += 1

        return chunks

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove multiple whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove multiple newlines
        text = re.sub(r'\n\n+', '\n\n', text)
        return text.strip()

    def _estimate_tokens(self, text: str) -> int:
        """Rough token count estimate"""
        # Approximate: 1 token ≈ 4 characters
        return len(text) // 4

    # ===================================
    # Entity Extraction (Basic)
    # ===================================

    def extract_entities(self, text: str) -> List[Dict]:
        """
        Basic entity extraction using patterns
        TODO: Enhance with NLP models (spaCy, etc.)
        """
        entities = []

        # Extract capitalized phrases (potential proper nouns)
        # Pattern: 2-4 consecutive capitalized words
        pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b'
        matches = re.findall(pattern, text)

        # Count occurrences
        entity_counts = {}
        for match in matches:
            entity_counts[match] = entity_counts.get(match, 0) + 1

        # Filter entities mentioned more than once
        for entity, count in entity_counts.items():
            if count > 1 and len(entity.split()) >= 2:
                entities.append({
                    'name': entity,
                    'type': 'unknown',  # TODO: Add type classification
                    'mentions': count
                })

        return entities

    # ===================================
    # Relationship Extraction (Basic)
    # ===================================

    def extract_relations(self, text: str, entities: List[Dict]) -> List[Dict]:
        """
        Basic relationship extraction
        TODO: Enhance with NLP models
        """
        relations = []
        entity_names = [e['name'] for e in entities]

        # Look for entities mentioned in same sentence
        sentences = re.split(r'[.!?]', text)

        for sentence in sentences:
            # Find which entities appear in this sentence
            entities_in_sentence = [e for e in entity_names if e in sentence]

            # Create relationships for entities in same sentence
            if len(entities_in_sentence) >= 2:
                for i, source in enumerate(entities_in_sentence):
                    for target in entities_in_sentence[i+1:]:
                        relations.append({
                            'source': source,
                            'target': target,
                            'type': 'mentioned_with',  # Generic relation type
                            'confidence': 0.5  # Low confidence for basic extraction
                        })

        return relations


# Singleton instance
_processor = None

def get_processor():
    """Get or create knowledge processor instance"""
    global _processor
    if _processor is None:
        _processor = KnowledgeProcessor()
    return _processor
